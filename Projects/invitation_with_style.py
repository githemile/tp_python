import docx

def create_invitations(guests_file, template_file, output_file):
    # Charger le fichier Word modèle
    doc = docx.Document(template_file)

    # Lire les noms des invités depuis le fichier texte
    with open(guests_file, 'r') as f:
        guests = f.readlines()

    for guest in guests:



        # Ajouter les informations supplémentaires
        doc.add_paragraph('You are cordially invited to a dinner party!', style='Normal')
        doc.add_paragraph('Date: April 1st, 2024', style='Normal')
        doc.add_paragraph('Time: 7:00 PM', style='Normal')
        doc.add_paragraph('Location: 123 Memory Lane', style='Normal')
        doc.add_paragraph('Please RSVP by March 25th', style='Normal')

        # Ajouter un saut de page après chaque invitation
        doc.add_page_break()

    # Sauvegarder le document final
    doc.save(output_file)

if __name__ == '__main__':
    guest_file = 'C:/Users/emile/PycharmProjects/LOVI_KOFFI_EMILE_PYTHON_TP_TEST/TP_PYTHON/files_projects/guests.txt'
    template_file = 'C:/Users/emile/PycharmProjects/LOVI_KOFFI_EMILE_PYTHON_TP_TEST/TP_PYTHON/files_projects/files/template.docx'
    output_file = 'C:/Users/emile/PycharmProjects/LOVI_KOFFI_EMILE_PYTHON_TP_TEST/TP_PYTHON/files_projects/files/invitation_style.docx'

    create_invitations(guest_file, template_file, output_file)
