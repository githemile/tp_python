import docx, readDocx

doc = docx.Document('C:/Users/emile/PycharmProjects/LOVI_KOFFI_EMILE_PYTHON_TP_TEST/TP_PYTHON/files_used_in_word/demo.docx')
len(doc.paragraphs)
doc.paragraphs[1].text
doc.paragraphs[1].text
len(doc.paragraphs[1].runs)
doc.paragraphs[1].runs[0].text


