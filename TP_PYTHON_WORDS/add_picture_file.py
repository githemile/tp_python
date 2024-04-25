import docx

doc = docx.Document()

# Ajouter une image au document
with open('C:/Users/emile/PycharmProjects/LOVI_KOFFI_EMILE_PYTHON_TP_TEST/TP_PYTHON/files_used_in_word/logo.png', 'rb') as img_file:
    doc.add_picture(img_file, width=docx.shared.Inches(1), height=docx.shared.Cm(4))

# Ajouter un saut de page après l'image
doc.paragraphs[0].runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)

# Ajouter des titres au document
doc.add_heading('Titre 0', level=0)
doc.add_heading('Titre 1', level=1)
doc.add_heading('Titre 2', level=2)
doc.add_heading('Titre 3', level=3)
doc.add_heading('Titre 4', level=4)

# Enregistrer le document
doc.save('C:/Users/emile/PycharmProjects/LOVI_KOFFI_EMILE_PYTHON_TP_TEST/TP_PYTHON/files_used_in_word/add_pictures.docx')
