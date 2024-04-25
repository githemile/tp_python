import docx

doc = docx.Document()


doc.add_heading('Titre 0', 0)
doc.add_heading('Titre 1', 1)
doc.add_heading('Titre 2', 2)
doc.add_heading('Titre 3', 3)
doc.add_heading('Titre 4', 4)
doc.save('titres.docx')




doc.add_paragraph('Bonjour, monde !')
paraObj1 = doc.add_paragraph('Ceci est un deuxième paragraphe.')
paraObj2 = doc.add_paragraph('Ceci est encore un autre paragraphe.')
#doc.save('C:/Users/emile/PycharmProjects/LOVI_KOFFI_EMILE_PYTHON_TP_TEST/TP_PYTHON/files_used_in_word/monde3Titre.docx')



doc.add_paragraph('Ceci est sur la première page !')
doc.paragraphs[0].runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)
doc.add_paragraph('Ceci est sur la deuxième page !')


doc.save('C:/Users/emile/PycharmProjects/LOVI_KOFFI_EMILE_PYTHON_TP_TEST/TP_PYTHON/files_used_in_word/deuxPages.docx')

