import docx
doc = docx.Document('C:/Users/emile/PycharmProjects/LOVI_KOFFI_EMILE_PYTHON_TP_TEST/TP_PYTHON/files_used_in_word/demo.docx')
print(doc.paragraphs[0].text)
print(doc.paragraphs[0].style) # L'identifiant exact peut être différent : _ParagraphStyle('Title') id: 3095631007984
doc.paragraphs[0].style = 'Normal'
print(doc.paragraphs[1].text)
print((doc.paragraphs[1].runs[0].text, doc.paragraphs[1].runs[1].text))

doc.paragraphs[1].runs[1].underline = True

doc.save('restyleds.docx')
